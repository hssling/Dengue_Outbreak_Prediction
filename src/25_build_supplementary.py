"""
25_build_supplementary.py
=========================
Build the supplementary information file and the resubmission cover letter for
the IJDSA revision.

Supplementary_IJDSA_R1.docx — numbered in order of first citation in the text:
  S1  Panel construction, eligibility and exclusions        (cited Section 2.3)
  S2  NCVBDC reconciliation, all overlapping state-years    (cited Section 2.4)
  S3  Per-state one-year-ahead forecast performance         (cited Section 2.7)
  S4  Complete forecasting and classification metrics       (cited Section 2.7)
  S5  Leakage experiment specification and results          (cited Section 2.11)
  S6  Audit of the previously submitted dataset             (cited Section 4.6)

Cover_Letter_IJDSA_R1.docx
  Short covering letter for the resubmission.
"""

import json
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REAL = "outputs/real"
OUTDIR = "MMI_submission_package"
SUPP = f"{OUTDIR}/Supplementary_IJDSA_R1.docx"
COVER = f"{OUTDIR}/Cover_Letter_IJDSA_R1.docx"
SUBMISSION_ID = "9abb84c1-1d65-44de-916a-fc708a733fe8"

M = json.load(open(f"{REAL}/real_metrics.json"))
P = json.load(open(f"{REAL}/panel_provenance.json"))


def styles(doc, body=10):
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(body)
    n.paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 12.5), ("Heading 2", 11)):
        s = doc.styles[name]
        s.font.name = "Times New Roman"; s.font.size = Pt(size)
        s.font.bold = True; s.font.color.rgb = RGBColor(0, 0, 0)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = sec.bottom_margin = Inches(0.8)


def shade(cell, hexcolor="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def df_table(doc, df, fontsize=8, max_rows=None):
    d = df if max_rows is None else df.head(max_rows)
    tbl = doc.add_table(rows=1, cols=len(d.columns))
    tbl.style = "Table Grid"
    for j, c in enumerate(d.columns):
        cell = tbl.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(str(c)); r.bold = True; r.font.size = Pt(fontsize)
        shade(cell)
    for _, row in d.iterrows():
        cells = tbl.add_row().cells
        for j, v in enumerate(row):
            cells[j].paragraphs[0].text = ""
            r = cells[j].paragraphs[0].add_run(
                f"{v:,.3f}" if isinstance(v, float) else f"{v:,}" if isinstance(v, int)
                else str(v))
            r.font.size = Pt(fontsize)
    return tbl


# --------------------------------------------------------------------------- #
def build_supplementary():
    doc = Document()
    styles(doc)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Supplementary Information"); r.bold = True; r.font.size = Pt(14)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Structural Determinants Without Short-Term Predictability: "
                  "A Leakage-Controlled Benchmark of State-Level Dengue Burden "
                  "in India, 2015–2024\n"
                  f"International Journal of Data Science and Analytics — "
                  f"Submission ID {SUBMISSION_ID}")
    r.italic = True; r.font.size = Pt(9.5)

    # Section order follows order of first citation in the manuscript.
    # ------------------------------- S1 -------------------------------- #
    doc.add_paragraph("Table S1. Panel construction, eligibility and exclusions",
                      style="Heading 1")
    rows = [
        ("OpenDengue India admin-1 annual records screened",
         f"{P['opendengue_india_admin1_year_rows']}"),
        ("Records from NCVBDC (MOH-IND) sources, 2015–2024", f"{P['moh_rows_in_window']}"),
        ("Source identifiers used", "; ".join(P["source_uuids"])),
        ("Duplicate state-year records requiring resolution",
         f"{P['duplicate_state_year_records_resolved']}"),
        ("Candidate states/UTs after name harmonisation", f"{P['states_all_admin1']}"),
        ("States/UTs with an observed value in all ten years",
         f"{P['states_complete_2015_2024']}"),
        ("Final analysis panel", f"{P['panel_state_years']} state-years"),
        ("States/UTs excluded", ", ".join(P["excluded_states"]) or "none"),
        ("Reason for exclusion",
         "; ".join(f"{k}: {v} missing year(s)"
                   for k, v in P["missing_years_by_excluded_state"].items()) or "n/a"),
        ("Genuine reported zero counts retained", f"{P['zero_count_state_years']} state-years"),
        ("Eligibility rule", P["exclusion_rule"]),
    ]
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    for k, v in rows:
        c = tbl.add_row().cells
        c[0].paragraphs[0].text = ""
        r = c[0].paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(9)
        c[0].width = Inches(2.6)
        c[1].paragraphs[0].text = ""
        r = c[1].paragraphs[0].add_run(v); r.font.size = Pt(9)
        c[1].width = Inches(4.2)

    doc.add_paragraph("Structural covariate coverage", style="Heading 2")
    doc.add_paragraph(
        "Structural covariates were unavailable for a small number of union "
        "territories. These values were median-imputed for modelling and flagged; "
        "all between-state associations in Section 3.4 exclude them (n = 32 states/UTs "
        "with fully observed covariates).")
    mi = P["structural_covariate_missing_states"]
    md = pd.DataFrame([(k, ", ".join(v) if v else "none") for k, v in mi.items()],
                      columns=["Covariate", "States/UTs without an observed value"])
    df_table(doc, md, fontsize=8.5)

    # ------------------------------- S2 -------------------------------- #
    doc.add_page_break()
    doc.add_paragraph("Table S2. Reconciliation of the analysis panel against the "
                      "NCVBDC state-wise annual bulletin", style="Heading 1")
    rec = pd.read_csv(f"{REAL}/provenance_reconciliation.csv")
    prov = M["dataset"]["provenance"]
    doc.add_paragraph(
        f"Every state-year in the analysis panel that overlaps the independently "
        f"published NCVBDC bulletin is listed below. Of {prov['overlapping_state_years']} "
        f"overlapping state-years, {prov['exact_matches']} match exactly "
        f"({prov['pct_exact']:.0f}%; median absolute difference "
        f"{prov['median_abs_pct_difference']:.1f}%). Overlap years: "
        f"{', '.join(str(y) for y in prov['overlap_years'])}.")
    r1 = rec[["state", "year", "cases", "ncvbdc_cases", "difference"]].copy()
    r1.columns = ["State/UT", "Year", "Analysis panel", "NCVBDC bulletin", "Difference"]
    r1["Year"] = r1["Year"].astype(int)
    for c in ["Analysis panel", "NCVBDC bulletin", "Difference"]:
        r1[c] = r1[c].astype(int)
    df_table(doc, r1, fontsize=7.5)

    # ------------------------------- S3 -------------------------------- #
    doc.add_page_break()
    doc.add_paragraph("Table S3. Per-state one-year-ahead forecast performance",
                      style="Heading 1")
    doc.add_paragraph(
        "Mean absolute error on the log(1+cases) scale for each state across the "
        "seven expanding-window forecast origins (2018–2024), for the multi-modal "
        "gradient-boosting model and the two leading baselines. Lower is better. "
        "The gradient-boosting model does not systematically outperform state "
        "climatology in any subgroup of states.")
    ps = pd.read_csv(f"{REAL}/per_state_performance.csv")
    piv = ps.pivot(index="state", columns="model", values="mae_log").reset_index()
    piv.columns.name = None
    piv = piv.rename(columns={"state": "State/UT",
                              "Gradient boosting (multi-modal)": "GBM (multi-modal)",
                              "State climatology (training mean)": "State climatology",
                              "Persistence (previous year)": "Persistence"})
    df_table(doc, piv.round(3), fontsize=7.5)

    # ------------------------------- S4 -------------------------------- #
    doc.add_page_break()
    doc.add_paragraph("Table S4. Complete forecasting and classification metrics",
                      style="Heading 1")
    f = pd.DataFrame(M["forecasting"]).T.reset_index().rename(columns={"index": "Model"})
    f = f.sort_values("r2_log", ascending=False)
    f = f[["Model", "r2_log", "r2_log_within_state", "mae_log", "mae_cases",
           "rmse_cases", "mean_per_year_r2", "sd_per_year_r2", "n_forecasts"]]
    f.columns = ["Model", "Pooled R²", "Within-state R²", "MAE (log)", "MAE (cases)",
                 "RMSE (cases)", "Mean per-origin R²", "SD per-origin R²", "n forecasts"]
    for c in f.columns[1:]:
        f[c] = pd.to_numeric(f[c]).round(3)
    df_table(doc, f, fontsize=8)

    doc.add_paragraph("Outbreak-year classification", style="Heading 2")
    C = M["classification"]
    cls = pd.DataFrame([
        ("AUC-ROC", round(C["auc_roc"], 3), "0.500 (chance)"),
        ("Average precision", round(C["auprc"], 3),
         f"{C['auprc_baseline_prevalence']:.3f} (prevalence)"),
        ("Brier score", round(C["brier"], 3), f"{C['brier_baseline']:.3f} (non-informative)"),
        ("Sensitivity", round(C["sensitivity"], 3), "—"),
        ("Specificity", round(C["specificity"], 3), "—"),
        ("Positive predictive value", round(C["ppv"], 3),
         f"{C['auprc_baseline_prevalence']:.3f} (prevalence)"),
        ("State-years evaluated", C["n"], "—"),
        ("Outbreak years", C["n_outbreak"], "—"),
    ], columns=["Metric", "Value", "Reference"])
    df_table(doc, cls, fontsize=8.5)

    doc.add_paragraph("Calibration (quintiles of predicted probability)", style="Heading 2")
    cal = pd.DataFrame(C["calibration"])
    cal.columns = ["Mean predicted probability", "Observed outbreak frequency", "n"]
    df_table(doc, cal.round(3), fontsize=8.5)

    doc.add_paragraph("Out-of-fold permutation importance", style="Heading 2")
    pi = M["permutation_importance"]
    imp = (pd.Series(pi["per_feature_mae_increase"]).sort_values(ascending=False)
           .reset_index())
    imp.columns = ["Predictor", "Increase in out-of-sample MAE when permuted"]
    imp.iloc[:, 1] = imp.iloc[:, 1].round(4)
    df_table(doc, imp, fontsize=8.5)
    doc.add_paragraph(
        f"Baseline out-of-fold MAE (log scale): {pi['baseline_mae_log']:.3f}. "
        + "; ".join(f"{k}: {v:.1f}%" for k, v in pi["grouped_percent"].items()) + ".")

    doc.add_paragraph("Between-state associations, all outcomes", style="Heading 2")
    rows = []
    for outcome, res in M["between_state"].items():
        if not isinstance(res, dict) or "rho" in res or "partial_rho" in res:
            continue
        for cov, v in res.items():
            rows.append((outcome.replace("_", " "), cov.replace("_", " "),
                         round(v["rho"], 3), f"{v['ci_low']:.2f} to {v['ci_high']:.2f}",
                         round(v["p"], 4), v["n"]))
    bs = pd.DataFrame(rows, columns=["Outcome", "Covariate", "Spearman ρ",
                                     "Bootstrap 95% CI", "p", "n"])
    df_table(doc, bs, fontsize=8)
    a = M["between_state"]["gdp_incidence_adjusted_for_health_index"]
    doc.add_paragraph(
        f"GDP per capita versus mean incidence, adjusted for the NITI Health Index as a "
        f"surveillance-capacity proxy: partial Spearman ρ = {a['partial_rho']:.3f}, "
        f"p = {a['p']:.5f}, n = {a['n']}.")

    # ------------------------------- S5 -------------------------------- #
    doc.add_page_break()
    doc.add_paragraph("Table S5. Leakage experiment", style="Heading 1")
    doc.add_paragraph(
        "Each specification below was run on the identical authenticated panel with "
        "identical model hyperparameters. Only the evaluation design differs, so the "
        "difference from the primary specification measures the optimistic bias "
        "attributable to that design defect alone.")
    L = M["leakage_experiment"]
    lk = pd.DataFrame([
        ("Leakage-free year-ahead (primary specification)",
         round(L["a_leakage_free_year_ahead"], 3), "—"),
        ("Non-temporal random-split cross-validation",
         round(L["c_random_split_cv"], 3),
         f"+{L['c_random_split_cv'] - L['a_leakage_free_year_ahead']:.3f}"),
        ("Unshifted three-year rolling mean (direct target leakage)",
         round(L["b_unshifted_rolling_mean"], 3),
         f"+{L['b_unshifted_rolling_mean'] - L['a_leakage_free_year_ahead']:.3f}"),
    ], columns=["Specification", "R² (log scale)", "Inflation"])
    df_table(doc, lk, fontsize=8.5)
    ob = pd.DataFrame([
        ("Expanding-window outbreak threshold (primary)",
         round(L["d_outbreak_auc_expanding_threshold"], 3), "—"),
        ("Full-panel outbreak threshold (look-ahead)",
         round(L["d_outbreak_auc_lookahead_threshold"], 3),
         f"+{L['d_outbreak_auc_lookahead_threshold'] - L['d_outbreak_auc_expanding_threshold']:.3f}"),
    ], columns=["Specification", "Outbreak AUC", "Inflation"])
    df_table(doc, ob, fontsize=8.5)

    # ------------------------------- S6 -------------------------------- #
    doc.add_page_break()
    doc.add_paragraph("Table S6. Audit of the previously submitted dataset, and "
                      "reproduction instructions", style="Heading 1")
    for para in [
        "Both reviewers questioned the provenance of the monthly panel used in the "
        "originally submitted manuscript. We audited it and confirmed their concern.",
        "The file used for that analysis, data/raw/dengue_climate_india.csv, is the "
        "output of a data-simulation routine retained in the project repository: "
        "create_synthetic_dengue_data() in src/01_fetch_data.py, executed under "
        "numpy.random.seed(42). Regenerating the routine and comparing it to the "
        "analysis file record by record gives a maximum absolute difference of 0.0 "
        "for cases, temperature, rainfall and humidity across all 1,800 rows. The "
        "routine hard-codes a seasonal multiplier of 3.0 for August–October, which is "
        "the source of the previously reported 74% seasonality contribution.",
        "A second file, data/raw/india_dengue_monthly_real.csv, was likewise not "
        "observed monthly data: src/01b_process_real_data.py distributes annual "
        "national totals across months using a fixed weight vector.",
        "In addition, the submitted feature matrix contained a three-month rolling "
        "case mean computed from the unshifted case series, so the predictor included "
        "the target month. The previously reported R² = 0.892 and AUC = 0.936 are "
        "therefore artefacts of simulated data combined with target leakage, and are "
        "formally withdrawn.",
        "No monthly admin-1 dengue series exists for India in the public domain. We "
        "scanned all three OpenDengue V1.3 extracts: India is represented by 672 "
        "admin-1 records at annual resolution and 12 monthly records, all national and "
        "all from 2024. The present study is specified at annual resolution for this "
        "reason.",
        "Reproduction. The three scripts below regenerate every number, table and "
        "figure in the manuscript from the archived public source files, in order and "
        "without manual intervention:",
    ]:
        doc.add_paragraph(para)
    for cmd, desc in [
        ("python src/20_build_real_panel.py",
         "streams the OpenDengue Spatial extract, builds the 350-state-year panel, "
         "and reconciles it against the NCVBDC bulletin"),
        ("python src/21_real_analysis.py",
         "runs the leakage-free benchmarks, classification, permutation importance, "
         "between-state analysis and leakage experiment"),
        ("python src/22_real_figures.py",
         "regenerates all seven manuscript figures"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(cmd); r.bold = True; r.font.name = "Consolas"; r.font.size = Pt(9)
        r2 = p.add_run(" — " + desc); r2.font.size = Pt(9.5)

    doc.save(SUPP)
    return SUPP


# --------------------------------------------------------------------------- #
def build_cover():
    doc = Document()
    styles(doc, body=10.5)
    doc.add_paragraph("Siddalingaiah H S, MD")
    doc.add_paragraph("Professor, Department of Community Medicine")
    doc.add_paragraph("Shridevi Institute of Medical Sciences and Research Hospital")
    doc.add_paragraph("Tumkur 572106, Karnataka, India")
    doc.add_paragraph("hssling@yahoo.com | ORCID 0000-0002-4771-8285")
    doc.add_paragraph()
    doc.add_paragraph("29 August 2026")
    doc.add_paragraph()
    doc.add_paragraph("Professor Xiangmin Zhou")
    doc.add_paragraph("Editor, International Journal of Data Science and Analytics")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Re: Revised submission — ID {SUBMISSION_ID}"); r.bold = True

    for para in [
        "Dear Professor Zhou,",

        "Please find enclosed our revised manuscript, together with a point-by-point "
        "response to both reviewers. We must draw your attention to a material change "
        "before you read further.",

        "Both reviewers questioned the provenance of the monthly dengue panel underlying "
        "our analysis. We investigated and confirmed that they were correct: the dataset "
        "used for the submitted analysis was a synthetic development dataset generated by "
        "a simulation routine in our own repository, not observed surveillance data. The "
        "submitted feature matrix also contained the unshifted rolling case statistic that "
        "Reviewer 1 identified as direct target leakage. The reported R² of 0.892 and AUC "
        "of 0.936 are artefacts of both problems and we withdraw them without reservation. "
        "This was our error, and we are sorry to have brought it to your reviewers.",

        "We have therefore replaced the empirical core of the paper rather than revising "
        "the text around it. The study is rebuilt on an authenticated panel of 350 "
        "state-years across 35 states and union territories (2015–2024), compiled from "
        "NCVBDC annual returns, with nothing interpolated or reconstructed. All 124 "
        "state-years that overlap the independently published NCVBDC bulletin reconcile "
        "exactly. As Reviewer 2 correctly observed, no continuous monthly admin-1 series "
        "exists for India, so annual resolution is the finest at which an authentic "
        "national panel can be built.",

        "The findings have changed. Under a strictly leakage-free design, no "
        "machine-learning model outperforms a state's own historical mean; within-state "
        "R² is at or below zero for every model; and outbreak-year classification performs "
        "at chance with worse-than-baseline calibration. We also quantify, on identical "
        "data, how far the original design defects inflate apparent performance — an "
        "unshifted rolling statistic alone raises R² from 0.551 to 0.759. What survives, "
        "and is stronger on real data than on the simulated data, is the between-state "
        "structural gradient: dengue incidence rises with urbanisation (ρ = 0.635) and "
        "GDP per capita (ρ = 0.609), the latter robust to adjustment for a "
        "surveillance-capacity proxy.",

        "The paper is consequently a rigorously benchmarked negative forecasting result "
        "combined with a robust spatial finding and a formal quantification of "
        "leakage-induced bias, rather than the predictive framework originally submitted. "
        "We believe this is of genuine value to the Journal's readership, given how "
        "frequently high accuracies are reported in this literature. We recognise, "
        "however, that it is a substantially different paper, and we would fully "
        "understand a decision to treat it as a new submission.",

        "Enclosed are: the revised manuscript in clean and tracked-changes form; the "
        "point-by-point response; and supplementary information containing the full "
        "reconciliation table, per-state performance, complete metrics, the leakage "
        "experiment and an audit of the previously submitted dataset.",

        "We are grateful to both reviewers. Their diagnosis was correct in every "
        "particular, and the paper is a far more honest piece of work as a result.",
    ]:
        doc.add_paragraph(para)

    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("Siddalingaiah H S, on behalf of all authors")
    doc.save(COVER)
    return COVER


if __name__ == "__main__":
    os.makedirs(OUTDIR, exist_ok=True)
    s = build_supplementary()
    c = build_cover()
    for f in (s, c):
        d = Document(f)
        print(f"Saved {f}  paragraphs={len(d.paragraphs)} tables={len(d.tables)}")
