"""
23_build_revision_docx.py
=========================
Build the IJDSA revision Word deliverables from reports/MANUSCRIPT_IJDSA_R1.md.

Produces two files with identical content and layout:

  Main_Manuscript_IJDSA_R1_clean.docx    submission copy
  Main_Manuscript_IJDSA_R1_tracked.docx  same text, with every change from the
                                         originally submitted manuscript marked
                                         as a genuine Word tracked revision
                                         (w:ins / w:del), so the editor can run
                                         Review > Accept/Reject normally.

Layout follows the original submission: a full-width banner (title, authors,
declarations, abstract, keywords) followed by a two-column body, with tables and
figures temporarily restored to full width.
"""

import difflib
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NEW_MD = "reports/MANUSCRIPT_IJDSA_R1.md"
OLD_MD = "reports/MANUSCRIPT_IJDSA.md"
FIGDIR = "outputs/figures_real"
OUTDIR = "MMI_submission_package/IJDSA_R1"
CLEAN = f"{OUTDIR}/Main_Manuscript_IJDSA_R1_clean.docx"
TRACKED = f"{OUTDIR}/Main_Manuscript_IJDSA_R1_tracked.docx"

REV_AUTHOR = "Siddalingaiah H S"
REV_DATE = "2026-08-29T00:00:00Z"

# Manuscript figure number -> source file. Numbering follows order of first
# mention in the text, as required by the journal, so the classification,
# importance and leakage panels are not in filename order.
FIGURES = {
    1: ("fig1_panel_provenance.png", 6.5),
    2: ("fig2_forecast_benchmarks.png", 6.1),
    3: ("fig4_classification.png", 6.6),
    4: ("fig7_permutation_importance.png", 6.5),
    5: ("fig5_between_state.png", 6.6),
    6: ("fig6_incidence_map.png", 4.0),
    7: ("fig3_leakage_experiment.png", 6.5),
}

_rev_id = [1000]


# --------------------------------------------------------------------------- #
# Layout helpers                                                              #
# --------------------------------------------------------------------------- #
def set_columns(section, n, space_twips=432):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(space_twips))


def base_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(10)
    n.paragraph_format.space_after = Pt(4)
    n.paragraph_format.line_spacing = 1.05
    for name, size in (("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5)):
        s = doc.styles[name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(3)


LIST_BULLET = re.compile(r"^[*-]\s+")
FIGLEGEND = re.compile(r"^\*\*Figure (\d)\.\*\*")


def split_list_marker(s):
    """Strip a leading markdown bullet. Returns (text, paragraph_style).

    This must happen BEFORE inline parsing: a leading "* " would otherwise be
    read as an unmatched italic delimiter and corrupt the rest of the line.
    """
    if LIST_BULLET.match(s):
        return LIST_BULLET.sub("", s, count=1), "List Bullet"
    return s, None


def parse_inline(text):
    """Parse markdown emphasis into (text, bold, italic, superscript) runs.

    Handles nesting (italic inside bold, as in "**... year *t*-1 ...**") by
    tracking each delimiter independently rather than matching outermost pairs.
    A delimiter with no partner later in the string is emitted literally, which
    keeps things like the author footnote marker in "^1,*^" intact.
    """
    out, buf = [], []
    bold = italic = sup = False

    def flush():
        if buf:
            out.append(("".join(buf), bold, italic, sup))
            del buf[:]

    i, n = 0, len(text)
    while i < n:
        if text.startswith("**", i) and (bold or "**" in text[i + 2:]):
            flush(); bold = not bold; i += 2; continue
        if text[i] == "*" and (italic or "*" in text[i + 1:]):
            flush(); italic = not italic; i += 1; continue
        if text[i] == "^" and (sup or "^" in text[i + 1:]):
            flush(); sup = not sup; i += 1; continue
        buf.append(text[i]); i += 1
    flush()
    return out


def _emit(run, piece):
    txt, bold, italic, sup = piece
    run.text = txt
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if sup:
        run.font.superscript = True


def add_runs(p, text):
    for piece in parse_inline(text):
        _emit(p.add_run(), piece)


# --------------------------------------------------------------------------- #
# Tracked-revision helpers                                                    #
# --------------------------------------------------------------------------- #
def _rev_el(tag):
    _rev_id[0] += 1
    e = OxmlElement(tag)
    e.set(qn("w:id"), str(_rev_id[0]))
    e.set(qn("w:author"), REV_AUTHOR)
    e.set(qn("w:date"), REV_DATE)
    return e


def add_tracked(p, text, mode):
    """Append `text` to paragraph `p` as normal / inserted / deleted content."""
    if not text:
        return
    if mode == "equal":
        add_runs(p, text)
        return
    wrapper = _rev_el("w:ins" if mode == "ins" else "w:del")
    for piece in parse_inline(text):
        run = p.add_run()
        _emit(run, piece)
        if mode == "del":
            # A deleted run must carry <w:delText>, not <w:t>.
            for t in run._r.findall(qn("w:t")):
                t.tag = qn("w:delText")
                t.set(qn("xml:space"), "preserve")
        wrapper.append(run._r)
    p._p.append(wrapper)


def _plain(text):
    """Drop emphasis markers.

    Word-level diffing splits on spaces, so a chunk can end up holding half of
    an emphasis pair ("*Healthy" without its closing "*"). Rendering the
    marked-up copy plain avoids stray literal asterisks; the clean copy keeps
    full formatting.
    """
    return re.sub(r"\*\*|\*|\^", "", text)


def word_diff(p, old, new):
    """Word-level tracked diff of two versions of one paragraph."""
    a, b = _plain(old).split(" "), _plain(new).split(" ")
    for op, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a, b).get_opcodes():
        if op == "equal":
            add_tracked(p, " ".join(a[i1:i2]) + " ", "equal")
        else:
            if i1 != i2:
                add_tracked(p, " ".join(a[i1:i2]) + " ", "del")
            if j1 != j2:
                add_tracked(p, " ".join(b[j1:j2]) + " ", "ins")


# --------------------------------------------------------------------------- #
# Content blocks                                                              #
# --------------------------------------------------------------------------- #
def _repeat_header(row):
    """Mark a row as a header so Word repeats it across page breaks."""
    trPr = row._tr.get_or_add_trPr()
    for tag in ("w:tblHeader", "w:cantSplit"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "true")
        trPr.append(el)


def _no_split(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def add_table(doc, lines, i):
    headers = [c.strip().replace("**", "") for c in lines[i].split("|")[1:-1]]
    # Column alignment comes from the markdown separator row (:--- / ---: / :-:)
    seps = [c.strip() for c in lines[i + 1].split("|")[1:-1]]
    aligns = []
    for spec in seps:
        if spec.startswith(":") and spec.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif spec.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            aligns.append(WD_ALIGN_PARAGRAPH.LEFT)

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = True
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = aligns[j] if j < len(aligns) else None
        r = cell.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    _repeat_header(tbl.rows[0])

    i += 2  # skip the |---| separator
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].split("|")[1:-1]]
        row = tbl.add_row()
        _no_split(row)
        for j, c in enumerate(cells):
            if j < len(row.cells):
                para = row.cells[j].paragraphs[0]
                para.text = ""
                # Numeric columns read better centred; the label column stays left.
                para.alignment = aligns[j] if j < len(aligns) else None
                add_runs(para, c)
                for rr in para.runs:
                    rr.font.size = Pt(9)
        i += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return i


def embed_figure(doc, number):
    """Place one figure image, immediately above its own legend."""
    entry = FIGURES.get(number)
    if not entry:
        return
    fname, width = entry
    path = os.path.join(FIGDIR, fname)
    if not os.path.exists(path):
        print(f"  [warn] missing figure {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True      # legend must not orphan
    p.add_run().add_picture(path, width=Inches(width))


def title_banner(doc, title):
    set_columns(doc.sections[0], 1)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title); r.bold = True; r.font.size = Pt(15)

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(a, "Siddalingaiah H S ^1,*^, Sowjanya D ^1^, Rangaswamy H V ^1^")
    aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = aff.add_run("1 Department of Community Medicine, Shridevi Institute of Medical "
                    "Sciences and Research Hospital, Tumkur 572106, Karnataka, India")
    r.font.size = Pt(9); r.italic = True
    cor = doc.add_paragraph(); cor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cor.add_run("*Corresponding author: Siddalingaiah H S, Professor of Community "
                    "Medicine. E-mail: hssling@yahoo.com; Tel.: +91-8941087719; "
                    "ORCID: 0000-0002-4771-8285")
    r.font.size = Pt(9)

    doc.add_paragraph("Statements and Declarations", style="Heading 2")
    decl = [
        ("Funding", "This research received no specific grant from any funding agency "
         "in the public, commercial, or not-for-profit sectors."),
        ("Competing Interests", "The authors declare no competing interests."),
        ("Ethics Approval", "This study used aggregated, anonymised, publicly available "
         "state-level data and did not involve individual human participants; ethics "
         "approval and individual consent were not applicable."),
        ("Consent to Participate", "Not applicable (no individual participant data)."),
        ("Consent for Publication", "Not applicable."),
        ("Availability of Data and Materials", "The authenticated state-year panel, the "
         "India admin-1 records extracted from the OpenDengue Spatial extract V1.3, the "
         "NCVBDC reconciliation table and all analysis outputs are openly available at "
         "https://github.com/hssling/Dengue_Outbreak_Prediction. A manifest records the "
         "origin, release version and SHA-256 checksum of every input file."),
        ("Code Availability", "All code for panel construction and reconciliation, "
         "leakage-free modelling, the benchmark suite, the leakage experiment and figure "
         "generation is available at the repository above, together with two automated "
         "verification suites. Every number and figure reported here is regenerated in a "
         "single pass; see Supplementary Table S6 for instructions."),
        ("Author Contributions", "Siddalingaiah H S: Conceptualization, Methodology, "
         "Software, Formal analysis, Visualization, Supervision, Validation, "
         "Writing - original draft. Sowjanya D: Writing - review & editing. "
         "Rangaswamy H V: Writing - review & editing."),
    ]
    for h, body in decl:
        p = doc.add_paragraph()
        rr = p.add_run(h + ". "); rr.bold = True; rr.font.size = Pt(9)
        r2 = p.add_run(body); r2.font.size = Pt(9)


# --------------------------------------------------------------------------- #
def read_md(path):
    with open(path, encoding="utf-8") as f:
        return [l.rstrip("\n") for l in f]


def prose_blocks(lines):
    """Body prose paragraphs only - used to align old and new for tracking."""
    out = []
    for l in lines:
        s = l.strip()
        if not s or s.startswith("|") or s.startswith("#") or set(s) <= {"-"}:
            continue
        # Strip bullets so these keys match the text used for tracked pairing.
        out.append(split_list_marker(s)[0])
    return out


def build(tracked=False):
    lines = read_md(NEW_MD)
    doc = Document()
    base_styles(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.75)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    title = next(l[2:] for l in lines if l.startswith("# "))
    title_banner(doc, title)

    # Map each new prose paragraph to its closest old counterpart so the
    # tracked copy shows edits rather than a wholesale replace where the
    # sentence actually survived.
    pairing = {}
    if tracked:
        old, new = prose_blocks(read_md(OLD_MD)), prose_blocks(lines)
        sm = difflib.SequenceMatcher(None, old, new)
        for op, i1, i2, j1, j2 in sm.get_opcodes():
            if op == "equal":
                for k in range(j2 - j1):
                    pairing[new[j1 + k]] = ("equal", old[i1 + k])
            elif op == "replace":
                for k in range(j1, j2):
                    # pair with the most similar old paragraph in this block
                    cand = max(old[i1:i2],
                               key=lambda o: difflib.SequenceMatcher(None, o, new[k]).ratio(),
                               default=None)
                    ratio = (difflib.SequenceMatcher(None, cand, new[k]).ratio()
                             if cand else 0)
                    pairing[new[k]] = ("edit", cand) if ratio > 0.55 else ("ins", None)
            elif op == "insert":
                for k in range(j1, j2):
                    pairing[new[k]] = ("ins", None)

    body_started = False
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if line.startswith("# ") or not s or set(s) <= {"-"}:
            i += 1
            continue

        if s.startswith("## 1. Introduction") and not body_started:
            doc.add_section(WD_SECTION.CONTINUOUS)
            ns = doc.sections[-1]
            ns.left_margin = ns.right_margin = Inches(0.75)
            set_columns(ns, 2)
            body_started = True

        if s.startswith("|"):
            if body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 1)
            i = add_table(doc, lines, i)
            if body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 2)
            continue

        if s.startswith("### "):
            doc.add_paragraph(s[4:], style="Heading 3"); i += 1; continue
        if s.startswith("## "):
            heading = s[3:]
            # Figures need the full page width, so the legends section runs
            # single-column; the reference list returns to two columns.
            if heading.startswith("7. Figure Legends") and body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 1)
            elif heading.startswith("8. References") and body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 2)
            doc.add_paragraph(heading, style="Heading 1")
            i += 1; continue

        # A figure legend carries its own image directly above it.
        mfig = FIGLEGEND.match(s)
        if mfig and body_started:
            embed_figure(doc, int(mfig.group(1)))

        body, list_style = split_list_marker(s)
        p = doc.add_paragraph(style=list_style) if list_style else doc.add_paragraph()
        s = body
        # A table caption must not be stranded at the foot of a column.
        if s.startswith("**Table "):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
        if tracked:
            mode, counterpart = pairing.get(s, ("ins", None))
            if mode == "equal":
                add_runs(p, s)
            elif mode == "edit":
                word_diff(p, counterpart, s)
            else:
                add_tracked(p, s, "ins")
        else:
            add_runs(p, s)
        i += 1

    # Record the removed material at the end of the tracked copy so that
    # rejecting every revision restores the original argument.
    if tracked:
        old_blocks = prose_blocks(read_md(OLD_MD))
        kept = {c for m, c in pairing.values() if c}
        removed = [o for o in old_blocks if o not in kept]
        if removed:
            doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 1)
            doc.add_paragraph("Text removed from the originally submitted manuscript",
                              style="Heading 1")
            for o in removed:
                add_tracked(doc.add_paragraph(), o, "del")

    out = TRACKED if tracked else CLEAN
    doc.save(out)
    d2 = Document(out)
    ncols = []
    for se in d2.sections:
        c = se._sectPr.find(qn("w:cols"))
        ncols.append(c.get(qn("w:num")) if c is not None else "1")
    print(f"Saved {out}\n  sections={len(d2.sections)} columns={ncols} "
          f"paragraphs={len(d2.paragraphs)} tables={len(d2.tables)}")


if __name__ == "__main__":
    os.makedirs(OUTDIR, exist_ok=True)
    build(tracked=False)
    build(tracked=True)
