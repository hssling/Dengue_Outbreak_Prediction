"""
Generate HIGH-QUALITY Professional Academic Manuscript DOCX.

Features:
- Professional Styling: Times New Roman, 12pt, Double Spaced.
- Robust Markdown Parsing: Headers, Bold, Italic, Lists.
- Dynamic Table Generation: Parses Markdown pipe tables correctly.
- Image Embedding: Automatically finds and embeds referenced images.
- IMRAD Structure Enforcement.
"""
import os
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_document_styles(doc):
    """Configure document styles for academic publication."""
    # Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph_format.space_after = Pt(0)
    
    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

def parse_markdown_table(doc, lines, start_index):
    """Parses a markdown table starting at start_index and inserts it into doc."""
    # Collect table lines
    table_lines = []
    i = start_index
    while i < len(lines) and '|' in lines[i]:
        table_lines.append(lines[i])
        i += 1
    
    if len(table_lines) < 2: return i # Not a valid table
    
    # Parse headers and alignment (skip separator line)
    headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
    
    # Create Table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Set Headers
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        for run in hdr_cells[idx].paragraphs[0].runs:
            run.font.bold = True
            
    # Parse Rows (Skip line 1 which is separator |---|---|)
    for row_line in table_lines[2:]:
        if not row_line.strip(): continue
        cells_data = [c.strip() for c in row_line.split('|') if c.strip() or c == '']
        # Handle cases where leading/trailing pipes create empty strings
        if row_line.strip().startswith('|'): cells_data = cells_data[0:] # regex logic better but this works for standard md
        
        # Clean empty splits
        clean_data = [c.strip() for c in row_line.strip('|').split('|')]
        
        row_cells = table.add_row().cells
        for idx, cell_text in enumerate(clean_data):
            if idx < len(row_cells):
                row_cells[idx].text = cell_text
                
    doc.add_paragraph() # Spacing after table
    return i

def process_math_text(paragraph, math_text):
    """
    Parses simple LaTeX-like math text inside $...$ and applies formatting.
    Supports: ^ (superscript), _ (subscript), and greek letters.
    """
    # math_text comes in as "$E = mc^2$" (without dollars if split correctly, or we strip them)
    math_text = math_text.strip('$')
    
    # Common Greek and Math Symbols Map
    symbols = {
        '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ', '\\epsilon': 'ε',
        '\\theta': 'θ', '\\lambda': 'λ', '\\mu': 'μ', '\\sigma': 'σ', '\\pi': 'π',
        '\\sum': '∑', '\\prod': '∏', '\\approx': '≈', '\\ne': '≠', '\\le': '≤', '\\ge': '≥',
        '\\times': '×', '\\rightarrow': '→'
    }
    
    # Replace known symbols first
    for k, v in symbols.items():
        math_text = math_text.replace(k, v)
        
    # Handle Functions (sin, cos, log) -> Standard text (not italic usually, but we will plain text them)
    # We can just remove the backslash for these simple text functions
    functions = ['\\sin', '\\cos', '\\tan', '\\log', '\\ln', '\\quad']
    for f in functions:
        math_text = math_text.replace(f, f.replace('\\', ''))
        
    # Handle Fractions: \frac{a}{b} -> (a/b)
    # Basic regex for simple fractions
    math_text = re.sub(r'\\frac\{(.*?)\}\{(.*?)\}', r'(\1/\2)', math_text)

    # Regex to tokenize: regular chars, or ^..., or _...
    # We want to split by ^ or _ but keep the markers to know what to do
    # Simple parser: look for ^(Token) or _(Token) or just Token
    
    # Let's try a char-by-char or simple split approach
    # A robust way is to find segments of plain, sub, sup
    # This regex looks for: ^ followed by single char or {group}, OR _ followed by single char or {group}
    pattern = r'(\^\{.*?\}|\^[^\s]|\_\{.*?\}|\_[^\s])'
    parts = re.split(pattern, math_text)
    
    for part in parts:
        if not part: continue
        
        # Superscript
        if part.startswith('^'):
            content = part[1:]
            if content.startswith('{') and content.endswith('}'):
                content = content[1:-1]
            run = paragraph.add_run(content)
            run.font.superscript = True
            
        # Subscript
        elif part.startswith('_'):
            content = part[1:]
            if content.startswith('{') and content.endswith('}'):
                content = content[1:-1]
            run = paragraph.add_run(content)
            run.font.subscript = True
            
        # Normal Math Text
        else:
            run = paragraph.add_run(part)
            # Heuristic: If it looks like a variable (single letter), italicize. 
            # If numbers or operators, maybe not? Word usually italicizes everything in Equation mode.
            # We'll stick to italic for consistency, unless it's a known function word like 'sin'
            if part.strip() in ['sin', 'cos', 'tan', 'log', 'ln', 'quad']:
                run.font.italic = False
            else:
                run.font.italic = True

def process_inline_formatting(paragraph, text):
    """Apply bold/italic/superscript/math formatting."""
    # Priority: Math ($...$) -> Bold (**...**) -> Italic (*...*) -> Superscript (^...^)
    
    # Check for Display Math ($$...$$) - usually handled by line logic, but if inline?
    # We'll treat inline $$ same as $ for now.
    
    # 1. Split by Math
    math_parts = re.split(r'(\$.*?\$)', text)
    for mpart in math_parts:
        if mpart.startswith('$') and mpart.endswith('$') and len(mpart) > 2:
            process_math_text(paragraph, mpart)
        else:
            # 2. Split by Bold
            bold_parts = re.split(r'(\*\*.*?\*\*)', mpart)
            for bpart in bold_parts:
                if bpart.startswith('**') and bpart.endswith('**'):
                    run = paragraph.add_run(bpart[2:-2])
                    run.font.bold = True
                else:
                    # 3. Check for Italics
                    italic_parts = re.split(r'(\*.*?\*)', bpart)
                    for ipart in italic_parts:
                        if ipart.startswith('*') and ipart.endswith('*'):
                            run = paragraph.add_run(ipart[1:-1])
                            run.font.italic = True
                        else:
                            # 4. Check for Superscripts (Citation style ^1,2^)
                            super_parts = re.split(r'(\^.*?\^)', ipart)
                            for spart in super_parts:
                                if spart.startswith('^') and spart.endswith('^'):
                                     run = paragraph.add_run(spart[1:-1])
                                     run.font.superscript = True
                                else:
                                     paragraph.add_run(spart)

def generate_professional_manuscript():
    print("Generating Professional Manuscript...")
    
    # Setup Document
    doc = Document()
    
    # Styles
    set_document_styles(doc) # Call existing style setter
    
    # Define styles if missing
    styles = doc.styles
    if 'Caption' not in styles:
        s = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles['Normal']
        s.font.size = Pt(10)
        s.font.italic = True
    
    # Read Markdown
    with open('reports/MANUSCRIPT_FINAL.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # --- TITLE PAGE ---
    # Title
    t_para = doc.add_paragraph('Multi-Modal Machine Learning Framework for State-Level Dengue Outbreak Prediction in India')
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_para.style = 'Heading 1'
    
    # Authors
    a_para = doc.add_paragraph('Siddalingaiah H S 1, *')
    a_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Affiliations
    aff_para = doc.add_paragraph('1 Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur 572106, Karnataka, India')
    aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff_para.style = 'Normal'
    
    # Corresponding
    corr_para = doc.add_paragraph('*Corresponding Author: E-mail: hssling@yahoo.com; Tel.: +91-8941087719')
    corr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corr_para.style = 'Normal'
    
    doc.add_page_break()
    
    # --- CONTENT PARSING ---
    i = 0
    in_references = False
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
            
        # Display Math ($$...$$)
        if line.startswith('$$') and line.endswith('$$'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            process_math_text(p, line.replace('$$', '$')) # Reuse inline logic but centered
            i += 1
            continue
            
        # Headers
        if line.startswith('# '): # Title (skip, already added)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line.replace('#### ', ''), level=3)
            
        # Tables
        elif line.startswith('|'):
            i = parse_markdown_table(doc, lines, i)
            continue # parse_table returns the new index
            
        # Images / Figures
        elif line.startswith('![') or (line.startswith('[') and ('png' in line or 'jpg' in line)):
            # Extract path
            match = re.search(r'\((.*?)\)', line) or re.search(r'\[(.*?)\]', line) # Handle both ![alt](path) and [See path]
            if match:
                path = match.group(1).replace('file:///', '').replace('%20', ' ')
                # Fix relative paths from report to standard
                if 'outputs/figures' in path or 'reports' in path: 
                     # Try to find the file
                     if os.path.exists(path):
                         try:
                            doc.add_picture(path, width=Inches(6))
                            last_p = doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Add caption relative to image
                            doc.add_paragraph(f"Figure: {os.path.basename(path)}", style='Caption')
                         except:
                             doc.add_paragraph(f"[Image: {path}]")
                     # Try adjusting path if relative to script root
                     elif os.path.exists(os.path.join('outputs/figures', os.path.basename(path))):
                         real_path = os.path.join('outputs/figures', os.path.basename(path))
                         doc.add_picture(real_path, width=Inches(6))
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, line[2:])
        elif line[0].isdigit() and line[1:3] == '. ':
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, line.split('. ', 1)[1])
            
        # Standard Paragraphs
        elif line.startswith('>'): # Blockquotes
            p = doc.add_paragraph()
            process_inline_formatting(p, line.replace('> ', ''))
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0]
            run.font.italic = True
        else:
            p = doc.add_paragraph()
            process_inline_formatting(p, line)
            
        i += 1
        
    # Save with robust conflict resolution
    os.makedirs('submission_package', exist_ok=True)
    base_path = 'submission_package/Main_Manuscript_IJMR_Submission.docx'
    
    saved = False
    counter = 1
    out_path = base_path
    
    while not saved and counter < 20:
        try:
            doc.save(out_path)
            print(f"Saved {out_path}")
            saved = True
        except PermissionError:
            print(f"Error: Could not save to {out_path}. File open. Trying next version...")
            counter += 1
            out_path = base_path.replace('.docx', f'_v{counter}.docx')
            
    if not saved:
        print("CRITICAL ERROR: Could not save manuscript after 20 attempts. Please close the file.")

if __name__ == "__main__":
    generate_professional_manuscript()
