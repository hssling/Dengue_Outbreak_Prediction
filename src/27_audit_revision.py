"""
27_audit_revision.py
====================
Adversarial audit of the revision. Unlike 26_verify_revision.py (which checks
that the manuscript agrees with the artefacts), this script tries to BREAK the
analysis:

  A. Future-information test. Corrupt every observation at or after the target
     year and confirm that not one predictor for that target year changes. If
     any feature moves, the pipeline leaks - the exact defect the reviewers
     found in the original submission.
  B. Baseline-honesty test. Confirm the winning baseline (state climatology) is
     computed from training years only and never sees the test year.
  C. Citation audit. Every in-text citation marker must resolve to a reference,
     and every reference must be cited at least once.
  D. Cross-document consistency. Numbers quoted in the response letter and
     supplementary must match the artefacts.

Exits non-zero on any failure.
"""

import importlib.util
import json
import os
import re
import sys

import numpy as np
import pandas as pd

spec = importlib.util.spec_from_file_location("ra", "src/21_real_analysis.py")
ra = importlib.util.module_from_spec(spec)
spec.loader.exec_module(ra)

PANEL = "data/processed/real_state_year_panel.csv"
MD = "reports/MANUSCRIPT_IJDSA_R1.md"
REAL = "outputs/real"
PKG = "MMI_submission_package/IJDSA_R1"
M = json.load(open(f"{REAL}/real_metrics.json"))

ok, fail = [], []


def check(label, cond, detail=""):
    (ok if cond else fail).append(f"{label}{(' — ' + detail) if detail else ''}")


# ------------------------------------------------------------------ A ----- #
def test_future_information():
    """Predictors for year T must not react to any data at year >= T."""
    base = pd.read_csv(PANEL)
    fb = ra.build_features(base)

    for T in (2019, 2021, 2024):
        corrupt = base.copy()
        mask = corrupt["year"] >= T
        # Violent corruption: anything leaking forward will move visibly.
        corrupt.loc[mask, "cases"] = corrupt.loc[mask, "cases"] * 1000 + 77777
        fc = ra.build_features(corrupt)

        a = fb[fb.year == T].sort_values("state").reset_index(drop=True)
        b = fc[fc.year == T].sort_values("state").reset_index(drop=True)

        moved = []
        for f in ra.FULL:
            va, vb = a[f].values.astype(float), b[f].values.astype(float)
            same = np.allclose(np.nan_to_num(va, nan=-999), np.nan_to_num(vb, nan=-999))
            if not same:
                moved.append(f)
        check(f"A. no predictor for {T} reacts to data at/after {T}",
              not moved, f"leaked: {moved}")

        # sanity: the corruption must actually be visible in the outcome,
        # otherwise the test proves nothing
        check(f"A. corruption is detectable in the {T} outcome",
              not np.allclose(a["cases"].values, b["cases"].values))


# ------------------------------------------------------------------ B ----- #
def test_baseline_honesty():
    """State climatology must be a training-window mean, not a full-panel mean."""
    d = ra.build_features(pd.read_csv(PANEL))
    ty = 2022
    tr = d[(d.year < ty) & d[ra.FULL].notna().all(axis=1)]
    te = d[(d.year == ty) & d[ra.FULL].notna().all(axis=1)]
    train_mean = tr.groupby("state")["cases"].mean()
    full_mean = d[d[ra.FULL].notna().all(axis=1)].groupby("state")["cases"].mean()

    recs, _ = ra.forecast_benchmarks(d)
    row = next(r for r in recs
               if r["year"] == ty and r["model"] == "State climatology (training mean)")
    pred = np.array(row["pred"])
    states = row["state"]
    from_train = np.log1p(np.array([train_mean.get(s, np.nan) for s in states]))
    from_full = np.log1p(np.array([full_mean.get(s, np.nan) for s in states]))

    check("B. climatology baseline uses TRAINING-window mean",
          np.allclose(pred, from_train, equal_nan=True))
    check("B. climatology baseline is NOT the full-panel mean",
          not np.allclose(pred, from_full, equal_nan=True))

    # And the leakage experiment must actually differ from the primary run.
    L = M["leakage_experiment"]
    check("B. leakage variants differ from the primary specification",
          L["b_unshifted_rolling_mean"] > L["a_leakage_free_year_ahead"] + 0.05
          and L["c_random_split_cv"] > L["a_leakage_free_year_ahead"] + 0.02,
          f"primary={L['a_leakage_free_year_ahead']:.3f}")


# ------------------------------------------------------------------ C ----- #
def test_citations():
    text = open(MD, encoding="utf-8").read()
    body = text.split("## 8. References")[0]
    reflist = text.split("## 8. References")[1]

    cited = set()
    for m in re.findall(r"\^([\d,\s]+)\^", body):
        for n in re.split(r"[,\s]+", m.strip()):
            if n.isdigit():
                cited.add(int(n))

    listed = set(int(m) for m in re.findall(r"^(\d+)\.\s", reflist, re.M))

    check("C. every in-text citation resolves to a reference",
          cited <= listed, f"dangling: {sorted(cited - listed)}")
    check("C. every listed reference is cited in the text",
          listed <= cited, f"uncited: {sorted(listed - cited)}")
    check("C. reference list is contiguous from 1",
          listed == set(range(1, max(listed) + 1)) if listed else False,
          f"listed: {sorted(listed)}")

    # Springer requires numbering by order of first mention.
    ref_order = []
    for m in re.findall(r"\^([\d,\s]+)\^", body):
        for n in re.split(r"[,\s]+", m.strip()):
            if n.isdigit() and int(n) not in ref_order:
                ref_order.append(int(n))
    check("C. references numbered in order of first mention",
          ref_order == sorted(ref_order), f"order: {ref_order}")


def test_sequential_display_items():
    """Tables, figures and supplementary items must be cited, and in order."""
    text = open(MD, encoding="utf-8").read()
    body = text.split("## 6. Tables")[0]

    def cited(pattern):
        seen = []
        for m in re.finditer(pattern, body):
            n = int(m.group(1))
            if n not in seen:
                seen.append(n)
        return seen

    figs = cited(r"\*\*Figure (\d)[A-C]?\*\*")
    tabs = cited(r"\*\*Table (\d)\*\*")
    supp = cited(r"\*\*Supplementary Table S(\d)\*\*")

    # Declared items
    legends = sorted({int(n) for n in
                      re.findall(r"\*\*Figure (\d)\.\*\*",
                                 text.split("## 7. Figure Legends")[1])})
    tbl_blocks = sorted({int(n) for n in re.findall(r"\*\*Table (\d)\.", text)})

    for label, order, declared in [("figures", figs, legends),
                                   ("tables", tabs, tbl_blocks),
                                   ("supplementary tables", supp, list(range(1, 7)))]:
        check(f"C. all {label} are cited in the text",
              set(order) == set(declared),
              f"cited {order} vs declared {declared}")
        check(f"C. {label} cited in sequential order",
              order == sorted(order), f"order: {order}")


# ------------------------------------------------------------------ D ----- #
def test_standalone_article():
    """The article and its supplement must read as a standalone paper.

    No submission-process metadata, no reviewer correspondence, no local file
    paths. That material belongs in the cover letter and the point-by-point
    response, which are correspondence and are checked separately.
    """
    from docx import Document

    text = open(MD, encoding="utf-8").read()

    # Phrases chosen to avoid false positives: "revised" legitimately appears in
    # Section 2.2 describing superseding ministry data releases.
    meta = ["originally submitted", "previously submitted", "this manuscript",
            "the reviewers", "Reviewer 1", "Reviewer 2", "referee",
            "point-by-point", "resubmission", "9abb84c1",
            "in the interest of full transparency"]
    paths = [r"src/\d", r"\.py\b", r"data/(raw|processed)", r"outputs/",
             r"\.csv\b", r"\.json\b", r"\.docx\b"]

    for m in meta:
        check(f"E. article free of submission metadata: '{m}'",
              m.lower() not in text.lower())
    for p in paths:
        hits = re.findall(p, text)
        check(f"E. article free of local file references: /{p}/",
              not hits, f"found {hits[:3]}")

    def doctext(path):
        d = Document(path)
        t = "\n".join(x.text for x in d.paragraphs)
        for tb in d.tables:
            for r in tb.rows:
                t += "\n" + "\t".join(c.text for c in r.cells)
        return t

    supp = doctext(f"{PKG}/Supplementary_IJDSA_R1.docx")
    for m in ["originally submitted", "previously submitted", "Reviewer",
              "9abb84c1", "Submission ID"]:
        check(f"E. supplement free of submission metadata: '{m}'",
              m.lower() not in supp.lower())

    # The correspondence documents must still carry the disclosure in full.
    for name in ["Cover_Letter_IJDSA_R1.docx", "Response_to_Reviewers_IJDSA_R1.docx"]:
        t = doctext(f"{PKG}/{name}").lower()
        check(f"E. {name} retains the provenance disclosure",
              "synthetic" in t and "withdraw" in t)


def test_cross_document():
    from docx import Document

    F, C, B = M["forecasting"], M["classification"], M["between_state"]["mean_incidence"]
    L = M["leakage_experiment"]
    key = {
        "0.637": F["State climatology (training mean)"]["r2_log"],
        "0.551": F["Gradient boosting (multi-modal)"]["r2_log"],
        "0.759": L["b_unshifted_rolling_mean"],
        "0.522": C["auc_roc"],
        "0.635": B["urban_pct"]["rho"],
        "0.609": B["gdp_pc"]["rho"],
    }
    for label, val in key.items():
        check(f"D. artefact rounds to {label}", f"{val:.3f}" == label,
              f"actual {val:.5f}")

    def doctext(p):
        d = Document(p)
        t = "\n".join(x.text for x in d.paragraphs)
        for tb in d.tables:
            for r in tb.rows:
                t += "\n" + "\t".join(c.text for c in r.cells)
        return t

    resp = doctext("MMI_submission_package/IJDSA_R1/Response_to_Reviewers_IJDSA_R1.docx")
    supp = doctext("MMI_submission_package/IJDSA_R1/Supplementary_IJDSA_R1.docx")
    cover = doctext("MMI_submission_package/IJDSA_R1/Cover_Letter_IJDSA_R1.docx")

    for name, t, must in [
        ("response letter", resp, ["0.637", "0.551", "0.759", "350", "124", "0.635"]),
        ("supplementary", supp, ["0.637", "0.551", "0.522", "124"]),
        ("cover letter", cover, ["350", "0.635", "0.609", "0.759"]),
    ]:
        missing = [m for m in must if m not in t]
        check(f"D. {name} quotes the artefact numbers", not missing,
              f"missing {missing}")

    # The disclosure belongs in the correspondence only. The supplement is part
    # of the article and is checked in test_standalone_article() for the
    # ABSENCE of this material.
    for name, t in [("response letter", resp), ("cover letter", cover)]:
        check(f"D. {name} discloses the synthetic-data error",
              "synthetic" in t.lower() and ("withdraw" in t.lower() or "withdrawn" in t.lower()))

    # Supplementary section labels referenced elsewhere must exist.
    for lab in ["Table S1", "Table S2", "S3.", "Table S4", "S5.", "S6."]:
        check(f"D. supplementary contains {lab}", lab in supp)


# --------------------------------------------------------------------------- #
if __name__ == "__main__":
    test_future_information()
    test_baseline_honesty()
    test_citations()
    test_sequential_display_items()
    test_standalone_article()
    test_cross_document()

    lines = [f"PASSED {len(ok)}", f"FAILED {len(fail)}", ""]
    if fail:
        lines.append("FAILURES:")
        lines += [f"  - {f}" for f in fail]
    else:
        lines.append("Adversarial audit clean.")
    lines.append("")
    lines.append("Checks run:")
    lines += [f"  ok  {o}" for o in ok]
    out = "\n".join(lines)
    os.makedirs(".work", exist_ok=True)
    open(".work/audit.log", "w", encoding="utf-8").write(out)
    print(f"PASSED {len(ok)} FAILED {len(fail)}")
    sys.exit(1 if fail else 0)
