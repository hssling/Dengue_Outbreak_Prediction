"""
13_build_double_column_docx.py
==============================
Build the IJDSA double-column Word manuscript from reports/MANUSCRIPT_IJDSA.md.

Layout (per Springer / IJDSA "double formatted column" request):
  * Full-width single-column banner: title, authors, affiliation, declarations,
    abstract and keywords.
  * Two-column body: Introduction -> References.
  * Wide tables and figures temporarily switch back to full width for
    readability, then resume two columns.
"""
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MD = "reports/MANUSCRIPT_IJDSA.md"
FIGDIR = "outputs/figures"
OUT = "MMI_submission_package/Main_Manuscript_IJDSA_double_column.docx"

FIGURES = {  # figure number -> (file, embed width inches) -- numbered by in-text mention order
    1: ("roc_curve.png", 3.0),
    2: ("feature_importance.png", 6.4),
    3: ("risk_vs_vulnerability.png", 3.4),
    4: ("india_risk_map.png", 3.2),
    "Box": ("practitioner_scorecard.png", 6.0),
}


# --------------------------------------------------------------------------- #
def set_columns(section, n, space_twips=432):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(space_twips))


def base_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(10)
    n.paragraph_format.space_after = Pt(4)
    n.paragraph_format.line_spacing = 1.05
    for name, size in (("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5)):
        s = doc.styles[name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(3)


def add_runs(p, text):
    """Bold **..**, italic *..*, superscript ^..^."""
    text = text.replace("R²", "R²").replace("²", "²")
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|\^[^\s|]+\^|\^[^\s|]+)", text)
    for t in tokens:
        if not t:
            continue
        if t.startswith("**") and t.endswith("**") and len(t) > 4:
            r = p.add_run(t[2:-2]); r.bold = True
        elif t.startswith("*") and t.endswith("*") and len(t) > 2:
            r = p.add_run(t[1:-1]); r.italic = True
        elif t.startswith("^"):
            r = p.add_run(t.strip("^")); r.font.superscript = True
        else:
            p.add_run(t)


def add_table(doc, lines, i):
    headers = [c.strip().replace("**", "") for c in lines[i].split("|")[1:-1]]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    i += 2  # skip separator
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].split("|")[1:-1]]
        row = tbl.add_row().cells
        for j, c in enumerate(cells):
            if j < len(row):
                row[j].paragraphs[0].text = ""
                add_runs(row[j].paragraphs[0], c)
                for rr in row[j].paragraphs[0].runs:
                    rr.font.size = Pt(9)
        i += 1
    return i


def embed_figure(doc, key, width):
    f = FIGURES.get(key)
    if not f:
        return
    path = os.path.join(FIGDIR, f[0])
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))


# --------------------------------------------------------------------------- #
def title_banner(doc, title):
    set_columns(doc.sections[0], 1)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title); r.bold = True; r.font.size = Pt(15)

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(a, "Siddalingaiah H S ^1,*^, Sowjanya D ^1^, Rangaswamy H V ^1^")
    aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = aff.add_run("1 Department of Community Medicine, Shridevi Institute of Medical "
                    "Sciences and Research Hospital, Tumkur 572106, Karnataka, India")
    r.font.size = Pt(9); r.italic = True
    cor = doc.add_paragraph(); cor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cor.add_run("*Corresponding author: Siddalingaiah H S, Professor of Community "
                    "Medicine. E-mail: hssling@yahoo.com; Tel.: +91-8941087719; "
                    "ORCID: 0000-0002-4771-8285")
    r.font.size = Pt(9)

    doc.add_paragraph("Statements and Declarations", style="Heading 2")
    decl = [
        ("Funding", "This research received no specific grant from any funding agency "
         "in the public, commercial, or not-for-profit sectors."),
        ("Competing Interests", "The authors declare no competing interests."),
        ("Ethics Approval", "This study used aggregated, anonymised, publicly available "
         "state-level data and did not involve individual human participants; ethics "
         "approval and individual consent were not applicable."),
        ("Consent to Participate", "Not applicable (no individual participant data)."),
        ("Consent for Publication", "Not applicable."),
        ("Availability of Data and Materials", "The processed dataset and all analysis "
         "code are openly available at the project repository: "
         "https://github.com/hssling/Dengue_Outbreak_Prediction"),
        ("Code Availability", "Code for preprocessing, modelling, validation and figure "
         "generation is available at the repository above; all reported statistics are "
         "regenerable from a single pipeline script."),
        ("Author Contributions", "Siddalingaiah H S: Conceptualization, Methodology, "
         "Software, Formal analysis, Visualization, Writing - original draft. "
         "Sowjanya D: Data curation, Investigation, Writing - review & editing. "
         "Rangaswamy H V: Supervision, Validation, Writing - review & editing."),
    ]
    for h, body in decl:
        p = doc.add_paragraph()
        rr = p.add_run(h + ". "); rr.bold = True; rr.font.size = Pt(9)
        r2 = p.add_run(body); r2.font.size = Pt(9)


def main():
    os.makedirs("MMI_submission_package", exist_ok=True)
    doc = Document()
    base_styles(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.75)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    with open(MD, encoding="utf-8") as f:
        lines = [l.rstrip("\n") for l in f]

    # Title line
    title = next(l[2:] for l in lines if l.startswith("# "))
    title_banner(doc, title)

    body_started = False
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if line.startswith("# ") or not s or set(s) <= {"-"}:
            i += 1
            continue

        # Abstract & keywords remain in the full-width banner; body starts at "## 1."
        if s.startswith("## 1. Introduction") and not body_started:
            doc.add_section(WD_SECTION.CONTINUOUS)
            ns = doc.sections[-1]
            ns.left_margin = ns.right_margin = Inches(0.75)
            set_columns(ns, 2)
            body_started = True

        # Tables -> full width
        if s.startswith("|"):
            if body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 1)
            i = add_table(doc, lines, i)
            if body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 2)
            continue

        # Headings
        if s.startswith("### "):
            doc.add_paragraph(s[4:], style="Heading 3"); i += 1; continue
        if s.startswith("## "):
            heading = s[3:]
            # figures: when we reach Figure Legends, embed images full width
            doc.add_paragraph(heading, style="Heading 1")
            if heading.startswith("7. Figure Legends") and body_started:
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 1)
                for key in (1, 2, 3, 4, "Box"):
                    embed_figure(doc, key, FIGURES[key][1])
                doc.add_section(WD_SECTION.CONTINUOUS); set_columns(doc.sections[-1], 2)
            i += 1; continue

        # Normal paragraph
        p = doc.add_paragraph()
        if s.startswith("**Keywords:**"):
            add_runs(p, s)
        else:
            add_runs(p, s)
        i += 1

    doc.save(OUT)
    print(f"Saved {OUT}")
    # quick audit
    d2 = Document(OUT)
    ncols = []
    for se in d2.sections:
        c = se._sectPr.find(qn("w:cols"))
        ncols.append(c.get(qn("w:num")) if c is not None else "1")
    print(f"Sections: {len(d2.sections)} | column pattern: {ncols} | "
          f"paragraphs: {len(d2.paragraphs)} | tables: {len(d2.tables)}")


if __name__ == "__main__":
    main()
