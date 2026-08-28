"""
24_build_response_letter.py
===========================
Build the point-by-point response to the editor and reviewers for IJDSA
submission 9abb84c1-1d65-44de-916a-fc708a733fe8, as a tabular Word document
plus a Markdown mirror.

Every response records (a) whether the point is accepted, (b) the concrete
action taken, and (c) where in the revised manuscript it can be verified.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTDIR = "MMI_submission_package/IJDSA_R1"
DOCX = f"{OUTDIR}/Response_to_Reviewers_IJDSA_R1.docx"
MD = f"{OUTDIR}/Response_to_Reviewers_IJDSA_R1.md"

SUBMISSION_ID = "9abb84c1-1d65-44de-916a-fc708a733fe8"

OPENING = [
    ("Dear Professor Zhou,"),
    ("Thank you for the opportunity to revise our manuscript, and for a review "
     "process that materially improved this work. We must begin with a "
     "disclosure."),
    ("Both reviewers questioned the provenance of the monthly panel underlying "
     "our analysis. We investigated this directly and confirmed that they were "
     "right. The file used for the submitted analysis, "
     "data/raw/dengue_climate_india.csv, is bit-identical to the output of a "
     "data-simulation routine retained in our own repository "
     "(src/01_fetch_data.py, create_synthetic_dengue_data, numpy seed 42). We "
     "regenerated the simulation and matched it to the analysis file record by "
     "record: the maximum absolute difference across all 1,800 rows was 0.0 for "
     "cases, temperature, rainfall and humidity. That routine hard-codes a "
     "seasonal multiplier of 3.0 for August-October, which is the origin of the "
     "“74% seasonality contribution” we reported. Reviewer 2's word for "
     "this - circular - is exactly correct. The submitted feature matrix also "
     "contained an unshifted three-month rolling case mean, precisely the target "
     "leakage Reviewer 1 identified at Section 2.4."),
    ("The reported R² of 0.892 and AUC of 0.936 were therefore artefacts of "
     "simulated data combined with target leakage. We withdraw them "
     "unreservedly. This was our error: a development dataset built early in the "
     "project to test the pipeline was never replaced with observed data before "
     "the analysis was finalised, and our internal checks failed to catch it."),
    ("Rather than revise the text around a compromised analysis, we have "
     "replaced the empirical core of the paper in its entirety, as Reviewer 1 "
     "indicated would be required. We first established what authentic data "
     "actually exist for India. Reviewer 2 is correct that no continuous monthly "
     "admin-1 series is available: we scanned all three OpenDengue V1.3 extracts "
     "and India has 672 admin-1 annual records but only 12 monthly records, all "
     "national and all from 2024. Annual state-level data are therefore the "
     "finest resolution at which a genuine Indian panel can be built, and we "
     "have rebuilt the study at that resolution: 350 state-years across 35 "
     "states and union territories, 2015-2024, drawn from NCVBDC returns "
     "archived by OpenDengue, with nothing interpolated or reconstructed. All "
     "124 state-years that overlap the independent NCVBDC bulletin match "
     "exactly (100%), which is the reconciliation table Reviewer 2 requested."),
    ("The scientific conclusion has changed accordingly, and we report it as we "
     "found it. Under a strictly leakage-free design, no machine-learning model "
     "beats a state's own historical mean, within-state R² is at or below zero "
     "for every model, and outbreak-year classification performs at chance with "
     "worse-than-baseline calibration. We also quantify, on identical data, how "
     "much the original defects inflate apparent performance. What survives - "
     "and is in fact stronger on real data than on the simulated data - is the "
     "between-state structural gradient: incidence rises with urbanisation "
     "(ρ = 0.635) and GDP per capita (ρ = 0.609), the latter robust to "
     "adjustment for a surveillance-capacity proxy."),
    ("We recognise that this is a substantially different paper from the one "
     "submitted, and we would understand entirely if you judge that it warrants "
     "assessment as a new submission. We would rather present an honest negative "
     "result on verified data than a strong positive result on data that cannot "
     "support it. We are genuinely grateful to both reviewers, whose diagnosis "
     "was correct in every particular."),
    ("Our point-by-point responses follow. Manuscript locations refer to the "
     "clean revised file; the tracked file shows every change as a Word "
     "revision."),
]

CLOSING = [
    ("Yours sincerely,"),
    ("Siddalingaiah H S, on behalf of all authors"),
    ("Department of Community Medicine, Shridevi Institute of Medical Sciences "
     "and Research Hospital, Tumkur, Karnataka, India"),
]

# (id, reviewer comment, response, where)
ITEMS = [
    # ---------------------------- REVIEWER 1 ------------------------------ #
    ("R1.1", "Data and Methods",
     "The provenance of the epidemiological and climatic data is inconsistent "
     "with the repository. The repository uses dengue_climate_india.csv, but "
     "src/01_fetch_data.py also provides a synthetic development dataset. The "
     "authors must identify the raw observed files and the executable workflow "
     "used to generate the presented findings. If the published conclusions "
     "were from the synthetic development dataset, the empirical analysis must "
     "be replicated using real surveillance data.",
     "Accepted in full; the reviewer is correct. We verified that "
     "dengue_climate_india.csv IS the synthetic development dataset: "
     "regenerating create_synthetic_dengue_data() with seed 42 reproduces it "
     "with a maximum absolute difference of 0.0 across all 1,800 rows. The "
     "entire empirical analysis has been replaced with an authenticated panel "
     "of 350 state-years (35 states/UTs, 2015-2024) built from NCVBDC annual "
     "returns archived in OpenDengue Spatial extract V1.3, restricted to "
     "MOH-IND source records. The executable workflow is now unambiguous and "
     "single-path: src/20_build_real_panel.py builds and reconciles the panel, "
     "src/21_real_analysis.py runs every model, and src/22_real_figures.py "
     "draws every figure. The simulation routine is retained in the repository "
     "only for historical transparency and is not on the analysis path.",
     "Section 2.2; Section 4.6; Figure 1"),

    ("R1.2", "Data and Methods",
     "The integration of NVBDCP and OpenDengue data is insufficiently "
     "documented. “Reconciled” does not clarify how the two sources were "
     "integrated. State the primary source, observations from each, treatment "
     "of missing or contradictory values, number substituted, and confirm no "
     "double counting. Interpolated or reconstructed data must be clearly "
     "distinguished from observed totals.",
     "Accepted. The ambiguity is removed because there is now a single source: "
     "NCVBDC annual state returns, accessed through the OpenDengue archive "
     "(these are the same numbers, not two sources being blended). Section 2.2 "
     "states the exact record filter used. Where a state-year appeared under "
     "more than one ministry release, the revised return supersedes the "
     "provisional one and no duplicate was summed - zero duplicate state-years "
     "required resolution in the final panel. Nothing is interpolated, "
     "imputed, substituted or augmented: any state-year without an observed "
     "value causes the state to be excluded rather than filled.",
     "Section 2.2; Section 2.3"),

    ("R1.3", "Data and Methods",
     "The three-month rolling case mean introduces direct target leakage: the "
     "code computes state_df['cases_rolling3'] = state_df['cases'].rolling(3)"
     ".mean() from the unshifted series, so the predictor includes the target "
     "month. Shift the case series before computing rolling statistics and "
     "regenerate all models, metrics, ablations, importances, tables and "
     "figures.",
     "Accepted; the reviewer's reading of the code is exactly right. In the "
     "revised pipeline every backward-looking statistic is formed by lagging "
     "the outcome series BEFORE any rolling or expanding operation, so no "
     "predictor can contain the target period. All results, tables and figures "
     "have been regenerated from scratch. We went further: because this defect "
     "is common in the literature, we quantified it. Reintroducing an "
     "unshifted three-year rolling mean on the identical real data inflates "
     "R² from 0.551 to 0.759 (+0.209). This is now reported as a substantive "
     "finding.",
     "Section 2.6; Section 2.11; Section 3.5; Table 4; Figure 3A"),

    ("R1.4", "Data and Methods",
     "The implemented task is not clearly a genuine month-ahead forecast. The "
     "model uses current meteorological factors for month t and case data "
     "including month t. Delineate the forecast origin, duration, reporting "
     "delays and predictor availability. A genuine design should use only "
     "predictors available by the end of month t to predict month t+1.",
     "Accepted. The forecast origin is now stated explicitly and enforced in "
     "code: using only information observable up to 31 December of year t-1, "
     "predict the reported total of year t. Every predictor is a lagged or "
     "expanding function of prior years plus time-invariant structural "
     "covariates. No contemporaneous covariate of any kind enters the model. "
     "We also removed the climate block entirely rather than reconstruct it "
     "(see R2.5), which eliminates the contemporaneous-climate problem at "
     "source. Because the data are annual, the horizon is one year, not four "
     "weeks, and we no longer make any four-week lead-time claim anywhere in "
     "the paper.",
     "Section 2.6; Section 2.5; Abstract"),

    ("R1.5", "Data and Methods",
     "The reported predictive performance must be recalculated; R², RMSE, MAE, "
     "AUC and sensitivity cannot be considered reliable until leakage and "
     "horizon problems are corrected. Table 1 also incorrectly interprets MAE "
     "as “Median absolute error”; MAE denotes Mean Absolute Error. The "
     "description of RMSE as a “month-ahead error” should be retained only if "
     "the design genuinely predicts the subsequent month.",
     "Accepted on every point. All metrics are recomputed under the "
     "leakage-free design and are materially different: pooled R² for gradient "
     "boosting is 0.551 (not 0.892) and outbreak AUC is 0.522 (not 0.936). MAE "
     "is now correctly labelled Mean Absolute Error throughout. RMSE is "
     "labelled simply as case-scale error; the “month-ahead” description is "
     "removed, since the horizon is one year.",
     "Table 1; Table 2; Section 3.2; Section 3.3"),

    ("R1.6", "Data and Methods",
     "The ablation findings do not demonstrate improved accuracy from "
     "multimodal integration (R² fell from 0.912 to 0.892). Explain why the "
     "full model was retained; do not imply that climate and vulnerability "
     "improved short-term forecasting. Repeat the ablation after correcting "
     "the rolling feature.",
     "Accepted, and the revised analysis makes the point far more sharply than "
     "the original ablation did. On real data no configuration - with or "
     "without structural covariates - beats state climatology, and the full "
     "multi-modal model (R² = 0.551) is essentially tied with simply repeating "
     "last year's count (0.550). We no longer retain a preferred multi-modal "
     "model or imply any accuracy benefit from it. The paper's conclusion is "
     "now that the multi-modal hypothesis was tested and not supported for "
     "temporal prediction at this resolution.",
     "Section 1.4; Section 3.2; Table 1; Section 4.1"),

    ("R1.7", "Data and Methods",
     "The state risk forecasts require clearer temporal specification and "
     "validation. The “forthcoming season” is undefined and the construction "
     "of future predictors is unexplained. The risk map should not be "
     "presented as a deployable forecasting product without independent "
     "validation.",
     "Accepted. The composite risk score, the risk-ranking table and the "
     "forecast-based risk map have all been withdrawn. Since the underlying "
     "model has no demonstrable forecasting skill, a forecast-derived risk "
     "product is not defensible and we do not offer one. The map is replaced "
     "by a purely descriptive choropleth of OBSERVED mean annual incidence "
     "over 2015-2024, explicitly labelled in both the figure title and legend "
     "as descriptive and not a forecast. Former Table 3 (highest-risk states) "
     "is deleted.",
     "Figure 6 and its legend; former Table 3 removed"),

    ("R1.8", "Introduction",
     "The conceptual interpretation of structural vulnerability requires "
     "refinement. Reported cases do not directly measure health-system "
     "weakness; the NITI Health Index is a broad composite rather than a "
     "dengue-specific vulnerability measure. Causal language should be "
     "moderated.",
     "Accepted. Section 1.3 now states directly that reported cases are not a "
     "measure of health-system weakness, that stronger systems may record more "
     "cases through better detection while weaker systems under-report, and "
     "that the Health Index is a broad composite of outcomes, governance and "
     "inputs rather than a dengue-specific measure. We state that the design "
     "can identify predictive and between-state associations but cannot "
     "establish a causal vulnerability-exposure mechanism. The term "
     "“vulnerability-exposure nexus” and all causal phrasing are removed. "
     "Notably, on real data the Health Index shows no association with "
     "incidence at all (ρ = 0.159, p = 0.38), which we report as such.",
     "Section 1.3; Section 3.4; Section 4.5"),

    ("R1.9", "Introduction",
     "The hypothesis of improved prediction is inconsistent with the "
     "subsequent ablation findings. The hypothesis may be presented as "
     "something tested, but the manuscript should acknowledge clearly that it "
     "was not supported.",
     "Accepted. Section 1.4 now presents the multi-modal hypothesis explicitly "
     "as the hypothesis the study tested, and states in the same paragraph "
     "that it was not supported and that we report that outcome rather than a "
     "favourable reformulation of it. The Abstract and Conclusions carry the "
     "same framing.",
     "Section 1.4; Abstract; Section 5"),

    ("R1.10", "State selection",
     "The criteria for selecting the 15 states require clarification. Define "
     "“continuous” and “good-quality” reporting with reproducible eligibility "
     "criteria; report how many missing months were permitted, how zero counts "
     "were distinguished from non-reporting, and list excluded states with "
     "reasons.",
     "Accepted. The subjective criteria are gone. Eligibility is now a single "
     "reproducible rule: a state or UT is included only if an observed NCVBDC "
     "annual total exists for every one of the ten study years; zero missing "
     "years are permitted and no missing value is imputed or carried forward. "
     "This yields 35 of 36 candidates. Exactly one state is excluded - West "
     "Bengal, for two missing years - and this is stated in the text. Six "
     "state-years carry genuine reported zeros in small union territories; "
     "these are distinguished from non-reporting (which causes exclusion) and "
     "retained as observed.",
     "Section 2.3; Supplementary Table S1"),

    ("R1.11", "State selection",
     "State selection may introduce surveillance-related selection bias, "
     "preferentially retaining states with stronger surveillance systems. This "
     "matters because the study then examines associations between reported "
     "burden and health-system capacity.",
     "Accepted and discussed. We note that requiring ten complete ANNUAL "
     "returns is a far weaker selection filter than requiring continuous "
     "monthly reporting, and that it retains 35 of 36 candidates rather than "
     "15, which substantially reduces but does not eliminate the concern. The "
     "residual bias is listed explicitly as a limitation, and we flag that "
     "ascertainment plausibly varies with the same structural covariates under "
     "study - which is precisely why we added the surveillance-capacity "
     "adjustment described in R2.7.",
     "Section 2.3; Section 4.5, limitations 4 and 6"),

    ("R1.12", "State selection",
     "Clarify whether the between-state outcome is cumulative burden or mean "
     "monthly cases, as terminology differs between the manuscript and the "
     "code.",
     "Accepted. The primary between-state outcome is now stated unambiguously "
     "as MEAN ANNUAL INCIDENCE PER 100,000, and the same quantity is used in "
     "the text, the table, the figure and the code. We additionally report "
     "mean annual cases and cumulative cases, and show that the count-scale "
     "results are dominated by population size (density leads at ρ = 0.558 "
     "while GDP falls to ρ = 0.056), which is why incidence is the appropriate "
     "outcome. All three are in the artefact file.",
     "Section 2.10; Section 3.4; Table 3; Figure 5"),

    ("R1.13", "Interpretation",
     "The feature-importance terminology requires correction: report grouped "
     "predictive or impurity-based importance rather than “Gini contribution” "
     "or causal contribution, and consider supporting the figure with "
     "out-of-sample permutation importance.",
     "Accepted. Impurity-based importance is no longer reported anywhere. It "
     "is replaced by out-of-fold permutation importance: for each forecast "
     "origin, each predictor is permuted 30 times in the held-out year and the "
     "increase in out-of-sample MAE recorded. Results are described as grouped "
     "predictive importance, and we state in the Methods that these are "
     "predictive importances, not causal contributions, and do not indicate "
     "direction of effect.",
     "Section 2.9; Section 3.3; Figure 7"),

    ("R1.14", "Interpretation",
     "The statement that Gradient Boosting tolerates missing data is "
     "technically inaccurate; scikit-learn's implementations do not natively "
     "accept missing values. Describe the actual imputation or deletion "
     "procedure.",
     "Accepted; the original statement was simply wrong and has been deleted. "
     "The revised Methods describe the actual procedure: states with any "
     "missing annual case value are excluded outright (no imputation of the "
     "outcome); structural covariate gaps, confined to three small union "
     "territories, are median-imputed AND flagged, and all between-state "
     "associations are computed only on states with observed covariates "
     "(n = 32).",
     "Section 2.3; Section 2.5; Section 2.10"),

    ("R1.15", "Interpretation",
     "The interpretation of feature importance should be moderated; the "
     "framework may be described as more interpretable, but impurity "
     "importance and ablation alone do not make it fully transparent.",
     "Accepted. All claims of full transparency are removed. We describe what "
     "the analysis supports and no more: permutation importance shows the "
     "models lean on summaries of a state's own level (previous-year count and "
     "expanding prior mean), which is consistent with the within-state result "
     "that they recover state identity rather than dynamics.",
     "Section 2.9; Section 3.3"),

    ("R1.16", "Results and figures",
     "The vulnerability results and Figure 3 should be aligned: the text "
     "discusses GDP per capita, Health Index and SECI while the figure shows "
     "the composite risk score against the Health Index. The figure should "
     "present the analysis discussed in the text.",
     "Accepted, and the circularity is removed at source. The composite risk "
     "score no longer exists (see R1.7), so a figure plotting it against one "
     "of its own inputs cannot arise. The new Figure 5 plots exactly the "
     "analysis discussed in the text: observed mean annual incidence against "
     "urban share, GDP per capita and the Health Index, each with Spearman ρ, "
     "bootstrap 95% CI and p-value matching Table 3 line for line.",
     "Figure 5; Table 3; Section 3.4"),

    ("R1.17", "Discussion",
     "The Discussion and Conclusions should be moderated after reanalysis. "
     "Claims of “genuine temporal validation” and strong accuracy cannot be "
     "sustained. The claim that near-unity AUCs collapse to ~0.94 under "
     "temporal blocking should be formally supported or removed. Avoid "
     "implying Gradient Boosting outperformed deep learning, since no such "
     "comparison was conducted. Limitations should acknowledge provenance, "
     "leakage, horizon uncertainty and absence of external validation.",
     "Accepted on every point. The Discussion now reports a negative result "
     "and makes no accuracy claim. The unsupported AUC assertion is replaced "
     "by a formal, reported experiment (Section 3.5, Table 4, Figure 3) "
     "measuring inflation from each design defect on identical data. The "
     "deep-learning comparison section is deleted in full, since we ran no "
     "such comparison. Limitations are rewritten as six numbered items "
     "covering temporal resolution, series length, the omitted climate block, "
     "passive-surveillance ascertainment, the ecological design and residual "
     "selection. A dedicated section formally withdraws the previously "
     "reported figures.",
     "Section 3.5; Section 4.1; Section 4.5; Section 4.6; Section 5"),

    # ---------------------------- REVIEWER 2 ------------------------------ #
    ("R2.1", "Data provenance",
     "Public NCVBDC sources report annual data, while OpenDengue does not "
     "offer continuous monthly admin-1 series for India. Specify the exact "
     "documents, bulletins or RTI records used to compile the May 2015 to "
     "December 2024 panel. Explicitly state if any monthly figures were "
     "interpolated or modelled from annual totals, as this would make findings "
     "like “74% seasonal contribution” circular.",
     "Accepted; the reviewer is correct on both counts and we are grateful for "
     "the precision of the diagnosis. We scanned all three OpenDengue V1.3 "
     "extracts: India has 672 admin-1 records at ANNUAL resolution and only 12 "
     "monthly records, all national and all from 2024. No continuous monthly "
     "admin-1 series exists. The monthly panel in the submitted manuscript was "
     "not compiled from any bulletin or RTI record - it was simulated, and the "
     "74% seasonality finding was indeed circular, reproducing a seasonal "
     "multiplier hard-coded in the simulation. We state this explicitly in the "
     "manuscript. The study is rebuilt at annual resolution, the finest at "
     "which authentic Indian state-level data exist, and the seasonality "
     "finding is withdrawn entirely.",
     "Section 2.2; Section 4.6; cover letter"),

    ("R2.2", "Data provenance",
     "Provide a reconciliation table comparing state-wise annual sums to "
     "official numbers. Discuss the plausibility of zero missing values and "
     "reporting disruptions during the COVID-19 period. Clarify the "
     "description of counts being “mostly near zero” alongside an implied mean "
     "of ~728 cases per state-month.",
     "Accepted. A full reconciliation against the independent NCVBDC "
     "state-wise annual bulletin is now provided as Supplementary Table S2 and "
     "plotted in Figure 1B: all 124 overlapping state-years match EXACTLY "
     "(100% agreement, median absolute difference 0.0%). On missingness, we no "
     "longer claim zero missing values - Section 2.3 reports that one state "
     "(West Bengal) fails the completeness rule and is excluded, and that six "
     "state-years are genuine reported zeros. On COVID-19, the authentic panel "
     "shows the disruption plainly: national reported cases fall from 157,315 "
     "in 2019 to 39,419 in 2020, and this is annotated in Figure 1A and "
     "discussed as a signature that any genuine Indian series must display. "
     "The contradictory “mostly near zero” description belonged to the "
     "simulated monthly data and is deleted; the observed annual distribution "
     "is described directly instead.",
     "Section 2.4; Section 3.1; Figure 1; Supplementary Table S2"),

    ("R2.3", "Model justification",
     "Table 2 shows adding 17 climate and vulnerability features degraded "
     "performance. Either reframe the paper around the strength of the "
     "8-feature seasonal-autoregressive model, or provide concrete empirical "
     "evidence justifying the 25-feature pipeline.",
     "Accepted, with an outcome neither option anticipated. On authentic data "
     "the parsimonious autoregressive model does not turn out to be strong "
     "either: gradient boosting on history alone reaches R² = 0.488, below "
     "persistence (0.550) and well below state climatology (0.637). We "
     "therefore reframe the paper around what the data actually support - a "
     "rigorously benchmarked negative forecasting result plus a robust "
     "between-state structural finding - rather than around either candidate "
     "model. No feature set is advocated.",
     "Section 3.2; Table 1; Section 4.1"),

    ("R2.4", "Baselines and calibration",
     "Compare model performance against standard benchmarks on identical "
     "splits: historical seasonal mean (climatology), persistence, seasonal "
     "naive, and standard SARIMA/GLM. Expand classification metrics beyond "
     "AUC-ROC to include Precision-Recall curves, Brier scores and calibration "
     "curves (per TRIPOD+AI).",
     "Accepted in full; this is now the backbone of the paper. Seven "
     "specifications are evaluated on identical expanding-window splits: "
     "persistence, state climatology, a global annual mean, negative-binomial "
     "GLMs with a log-population offset (history and full), and gradient "
     "boosting (history and multi-modal). Table 1 reports all of them. "
     "Seasonal-naive and SARIMA are not applicable at annual resolution, where "
     "there is no within-year seasonal cycle to model; persistence and "
     "climatology are their annual analogues and both are included. "
     "Classification now reports AUC-ROC, precision-recall with the prevalence "
     "reference, Brier score against a non-informative baseline, and a "
     "quintile calibration curve, following TRIPOD+AI (now cited as reference "
     "13). Calibration is poor: Brier 0.326 against a baseline of 0.248.",
     "Section 2.7; Section 2.8; Table 1; Table 2; Figure 2; Figure 4"),

    ("R2.5", "Spatial confounding",
     "Pooling data across states with vastly different baseline burdens risks "
     "inflating R² simply because the model learns state means rather than "
     "temporal variations. Report per-state performance, within-state "
     "(demeaned) metrics, and explicitly benchmark against a “state mean only” "
     "baseline.",
     "Accepted, and this proved to be the single most consequential comment in "
     "the review. All three requested analyses are implemented, and the "
     "reviewer's hypothesis is confirmed. The “state mean only” baseline "
     "(state climatology) is not merely competitive - it WINS, with the "
     "highest pooled R² of any specification (0.637). Within-state R², computed "
     "after removing each state's training-window mean from both observed and "
     "predicted values, is at or below zero for every model (best -0.002; "
     "gradient boosting -0.239). Pooled skill is therefore entirely "
     "between-state. Per-state performance is provided as Supplementary Table "
     "S3. This finding now anchors Section 3.3 and the Discussion.",
     "Section 2.7; Section 3.3; Table 1; Figure 2; Supplementary Table S3"),

    ("R2.6", "Leakage safeguards",
     "Lead time: contemporaneous climate features conflict with the claimed "
     "4-week lead time; restrict environmental variables to lags >= 1 or "
     "clarify forecast inputs. Look-ahead bias: calculate outbreak thresholds "
     "via an expanding window rather than the full 10-year period, and select "
     "classification thresholds strictly within training folds.",
     "Accepted on all three points. On lead time, we removed the climate block "
     "entirely rather than lag a reconstructed series: state-resolution "
     "meteorological fields covering the full window were not available to us, "
     "and we prefer to omit a block and declare it as a limitation than to "
     "reconstruct one. No contemporaneous covariate of any kind now enters the "
     "model. On look-ahead bias, outbreak thresholds are recomputed inside "
     "each expanding training window, and the decision cut-off is selected by "
     "maximising Youden's J on TRAINING predictions only. We also quantified "
     "the reviewer's concern: using a full-panel threshold instead inflates "
     "AUC from 0.522 to 0.544.",
     "Section 2.5; Section 2.6; Section 2.8; Section 3.5; Table 4; Figure 3B"),

    ("R2.7", "Interpretability and claims",
     "Use permutation importance or SHAP on out-of-fold data instead of Gini "
     "importance. Reframe the GDP correlation cautiously to account for "
     "surveillance intensity confounding in wealthier states. Define the "
     "formula, horizon and cut-offs for the 0-100 risk score and resolve the "
     "circularity in Figure 3. Ensure the Box 1 field tool matches the "
     "validated 0-6 instrument rather than an unvalidated scale with "
     "unmeasured entomological metrics.",
     "Accepted on all four. (i) Gini importance is replaced by out-of-fold "
     "permutation importance (30 permutations per predictor per origin). "
     "(ii) The GDP association is reframed with explicit attention to "
     "surveillance intensity: we now report incidence rather than counts, give "
     "bootstrap confidence intervals, and add a partial Spearman correlation "
     "adjusting for the Health Index as a detection-capacity proxy "
     "(ρ = 0.663, p < 0.001). We state plainly in the limitations that this "
     "mitigates but cannot eliminate the confounding, since the Health Index "
     "is an imperfect proxy for dengue-specific diagnostic capacity. (iii) The "
     "0-100 composite risk score is withdrawn rather than redefined, which "
     "dissolves the Figure 3 circularity (see R1.7 and R1.16). (iv) Box 1 and "
     "the field scorecard are withdrawn in full: the instrument was validated "
     "only against the discredited synthetic model's own predictions, so its "
     "reported ρ = 0.76 carried no independent meaning. We do not offer a "
     "replacement field tool, as the present analysis provides no validated "
     "basis for one.",
     "Section 2.9; Section 2.10; Section 3.4; Section 4.5; Box 1 and former "
     "Table 3 removed"),

    ("R2.8", "Minor",
     "Correct the Table 1 MAE label (Mean, not Median); clarify IMD data "
     "sources for relative humidity; update the 2009 WHO reference.",
     "All three corrected. MAE is labelled Mean Absolute Error throughout. The "
     "IMD humidity question is resolved by removal: no climate variables are "
     "used, so no IMD source needs to be characterised, and the unsupported "
     "IMD attribution is deleted from the Methods. The 2009 WHO guidelines "
     "citation is replaced with the current WHO Dengue and Severe Dengue fact "
     "sheet (2024, accessed 20 August 2026). We also added references for "
     "TRIPOD+AI, for leakage in machine-learning-based science, for evaluation "
     "metrics in medical AI, and for COVID-19 disruption of Indian dengue "
     "surveillance.",
     "Table 1; Section 2.5; References 1, 11, 13, 15, 16"),
]


# --------------------------------------------------------------------------- #
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def build_docx():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(10)
    n.paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 13), ("Heading 2", 11)):
        s = doc.styles[name]
        s.font.name = "Times New Roman"; s.font.size = Pt(size)
        s.font.bold = True; s.font.color.rgb = RGBColor(0, 0, 0)

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Response to the Editor and Reviewers"); r.bold = True; r.font.size = Pt(14)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Apparent Machine Learning Skill in Indian Dengue Forecasting "
                    "Vanishes Under Leakage-Free Within-State Evaluation\n"
                    "(submitted as “Multi-Modal Machine Learning Framework for "
                    "State-Level Dengue Outbreak Prediction in India”)\n"
                    f"International Journal of Data Science and Analytics — "
                    f"Submission ID {SUBMISSION_ID}")
    r.italic = True; r.font.size = Pt(9.5)

    doc.add_paragraph("Letter to the Editor", style="Heading 1")
    for para in OPENING:
        doc.add_paragraph(para)

    doc.add_paragraph("Point-by-Point Responses", style="Heading 1")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    widths = [Inches(0.55), Inches(2.7), Inches(3.6), Inches(1.3)]
    for j, head in enumerate(["#", "Reviewer comment", "Author response and action taken",
                              "Where in revised MS"]):
        c = tbl.rows[0].cells[j]
        c.paragraphs[0].text = ""
        rr = c.paragraphs[0].add_run(head); rr.bold = True; rr.font.size = Pt(9.5)
        shade(c, "D9E2F3")
        c.width = widths[j]

    current_reviewer = None
    for rid, section, comment, response, where in ITEMS:
        rev = "Reviewer 1" if rid.startswith("R1") else "Reviewer 2"
        if rev != current_reviewer:
            row = tbl.add_row().cells
            row[0].merge(row[3])
            m = tbl.rows[-1].cells[0]
            m.paragraphs[0].text = ""
            rr = m.paragraphs[0].add_run(rev); rr.bold = True; rr.font.size = Pt(10)
            shade(m, "EDEDED")
            current_reviewer = rev

        row = tbl.add_row().cells
        for j, txt in enumerate([rid, f"[{section}] {comment}", response, where]):
            row[j].paragraphs[0].text = ""
            rr = row[j].paragraphs[0].add_run(txt)
            rr.font.size = Pt(8.5)
            if j == 0:
                rr.bold = True
            row[j].width = widths[j]

    doc.add_paragraph()
    doc.add_paragraph("Summary of Changes", style="Heading 1")
    for line in [
        "The empirical analysis has been replaced in full. The synthetic development "
        "dataset is off the analysis path; the study now uses 350 authenticated "
        "state-years across 35 states and union territories (2015-2024), with all 124 "
        "overlapping state-years reconciling exactly against the official NCVBDC bulletin.",
        "The forecasting design is now strictly leakage-free: every predictor is "
        "observable before the forecast origin, outbreak thresholds and decision cut-offs "
        "are derived inside training windows, and no contemporaneous covariate is used.",
        "Seven models are benchmarked on identical splits, with pooled and within-state "
        "skill, precision-recall, Brier score and calibration reported.",
        "The headline results have changed: no model beats state climatology, within-state "
        "R² is at or below zero throughout, and outbreak classification is at chance. The "
        "previously reported R² = 0.892 and AUC = 0.936 are formally withdrawn.",
        "A new leakage experiment quantifies the bias from an unshifted rolling statistic "
        "(+0.209 R²), non-temporal validation (+0.075 R²) and look-ahead thresholds "
        "(+0.022 AUC).",
        "The between-state structural finding is retained, strengthened and correctly "
        "framed: incidence rises with urbanisation and GDP per capita, robust to "
        "adjustment for a surveillance-capacity proxy; the Health Index shows no "
        "association.",
        "The composite risk score, risk-ranking table, forecast risk map and field "
        "scorecard are withdrawn, along with the deep-learning comparison and all "
        "seasonality claims.",
        "The title has changed to match the revised content, from “Multi-Modal Machine "
        "Learning Framework for State-Level Dengue Outbreak Prediction in India” to "
        "“Apparent Machine Learning Skill in Indian Dengue Forecasting Vanishes Under "
        "Leakage-Free Within-State Evaluation”. The original title promised a "
        "predictive framework the authenticated data do not support.",
        "Both a clean and a tracked-changes manuscript file are supplied.",
    ]:
        p = doc.add_paragraph(line, style="List Bullet")
        for r_ in p.runs:
            r_.font.size = Pt(9.5)

    doc.add_paragraph()
    for para in CLOSING:
        doc.add_paragraph(para)

    doc.save(DOCX)
    return doc


def build_md():
    L = ["# Response to the Editor and Reviewers", "",
         "**Apparent Machine Learning Skill in Indian Dengue Forecasting Vanishes "
         "Under Leakage-Free Within-State Evaluation**", "",
         "*(submitted as “Multi-Modal Machine Learning Framework for State-Level "
         "Dengue Outbreak Prediction in India”)*", "",
         f"*International Journal of Data Science and Analytics — Submission ID "
         f"{SUBMISSION_ID}*", "", "## Letter to the Editor", ""]
    L += [p + "\n" for p in OPENING]
    L += ["", "## Point-by-Point Responses", "",
          "| # | Reviewer comment | Author response and action taken | Where in revised MS |",
          "| :--- | :--- | :--- | :--- |"]
    cur = None
    for rid, section, comment, response, where in ITEMS:
        rev = "Reviewer 1" if rid.startswith("R1") else "Reviewer 2"
        if rev != cur:
            L.append(f"| **{rev}** | | | |")
            cur = rev
        c = comment.replace("|", "/")
        r = response.replace("|", "/")
        L.append(f"| **{rid}** | [{section}] {c} | {r} | {where} |")
    L += ["", "## Closing", ""] + [p for p in CLOSING]
    with open(MD, "w", encoding="utf-8") as f:
        f.write("\n".join(L))


if __name__ == "__main__":
    os.makedirs(OUTDIR, exist_ok=True)
    build_docx()
    build_md()
    d = Document(DOCX)
    print(f"Saved {DOCX}")
    print(f"  paragraphs={len(d.paragraphs)} tables={len(d.tables)} "
          f"rows={len(d.tables[0].rows)} response_items={len(ITEMS)}")
    print(f"Saved {MD}")
